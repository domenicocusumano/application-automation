"""
Resume Pipeline — Dom Cusumano
Takes top jobs from job_scraper.py output, fetches each job description,
scores fit with Claude, and builds a tailored .docx resume for any role scoring 7+.
Uploads finished resumes to Google Drive automatically.

SETUP:
  pip3 install anthropic gspread google-auth google-api-python-client playwright
  npm install docx   (in the project directory, NOT -g)

USAGE:
  python3 resume_pipeline.py              # standalone with pasted jobs
  python3 job_scraper.py --resume         # integrated: scraper feeds into pipeline
"""

import warnings
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", message=".*urllib3.*")

import os
import re
import json
import time
import base64
import subprocess
import tempfile
import anthropic
from dotenv import load_dotenv

from playwright.sync_api import sync_playwright
from google.oauth2.service_account import Credentials as ServiceAccountCredentials
from google.oauth2.credentials import Credentials as OAuthCredentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

load_dotenv()

# ── CONFIG ─────────────────────────────────────────────────────────────────────

ANTHROPIC_API_KEY   = os.getenv("ANTHROPIC_API_KEY",  "YOUR_ANTHROPIC_API_KEY")
GOOGLE_CREDS_FILE   = os.getenv("GOOGLE_CREDS_FILE",   "google_credentials.json")
GDRIVE_FOLDER_ID    = os.getenv("GDRIVE_FOLDER_ID",    "")

# OAuth 2.0 credentials for Drive uploads (personal account — service accounts
# have no storage quota and cannot upload files to personal My Drive).
# oauth_credentials.json  → downloaded from Google Cloud Console (Desktop app type)
# gdrive_token.json       → auto-created on first run, reused on every run after
GDRIVE_OAUTH_CREDS  = os.getenv("GDRIVE_OAUTH_CREDS",  "oauth_credentials.json")
GDRIVE_TOKEN_FILE   = os.getenv("GDRIVE_TOKEN_FILE",   "gdrive_token.json")
GDRIVE_OAUTH_SCOPES = ["https://www.googleapis.com/auth/drive.file"]

SCRIPT_DIR          = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE         = os.path.join(SCRIPT_DIR, "config.json")


def _load_pipeline_config():
    """Reads scoring/filter settings and applicant info from config.json."""
    defaults = {
        "score_threshold":     6.5,
        "salary_minimum":      0,
        "preferred_locations": ["remote"],
        "resume_prefix":       "My_Resume",
        "gdrive_folder_id":    "",
    }
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            cfg = json.load(f)
            result = {k: cfg.get(k, v) for k, v in defaults.items()}
            # gdrive_folder_id can also come from the top-level env var
            if not result["gdrive_folder_id"]:
                result["gdrive_folder_id"] = GDRIVE_FOLDER_ID
            return result
    except (FileNotFoundError, ValueError):
        return defaults


score_threshold = _load_pipeline_config()["score_threshold"]

# ── DOM'S BACKGROUND ───────────────────────────────────────────────────────────

def _load_prompt_file(filename):
    path = os.path.join(SCRIPT_DIR, filename)
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

BACKGROUND_PROMPT = _load_prompt_file("background_prompt.txt")


# ── GOOGLE SERVICES ───────────────────────────────────────────────────────────

def _get_sheet_id() -> str:
    """Returns the Google Sheet ID from config.json (preferred) or .env."""
    try:
        import json as _json, re as _re
        cfg_path = Path(__file__).parent / "config.json"
        cfg = _json.loads(cfg_path.read_text())
        url_or_id = cfg.get("google_sheet_url", "").strip()
        if url_or_id:
            m = _re.search(r'/spreadsheets/d/([a-zA-Z0-9_-]+)', url_or_id)
            if m:
                return m.group(1)
            if _re.match(r'^[a-zA-Z0-9_-]+$', url_or_id):
                return url_or_id
    except Exception:
        pass
    return os.getenv("GOOGLE_SHEET_ID", "1IhPY7ukaZh5CAV2ZILWLwEDnrg_KXJcs-blNIqKuCpU")


SHEET_ID = _get_sheet_id()


def get_drive_service():
    """
    Returns an authenticated Drive service using OAuth 2.0 (personal account).
    Service accounts have no storage quota and cannot upload to personal My Drive.

    First run: opens a browser tab to authorize — takes ~10 seconds.
    All subsequent runs: uses the saved token silently, no browser needed.

    Setup (one-time):
      1. Google Cloud Console → APIs & Services → Credentials
      2. Create Credentials → OAuth 2.0 Client ID → Desktop app
      3. Download JSON → save as oauth_credentials.json in this directory
      4. Run the pipeline — a browser tab opens, click Allow, done.
    """
    token_path  = os.path.join(SCRIPT_DIR, GDRIVE_TOKEN_FILE)
    oauth_path  = os.path.join(SCRIPT_DIR, GDRIVE_OAUTH_CREDS)

    creds = None

    # Load saved token if it exists
    if os.path.exists(token_path):
        creds = OAuthCredentials.from_authorized_user_file(token_path, GDRIVE_OAUTH_SCOPES)

    # Refresh or run the full OAuth flow if needed
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists(oauth_path):
                raise FileNotFoundError(
                    f"OAuth credentials file not found: {oauth_path}\n"
                    "See the setup instructions in get_drive_service() above."
                )
            flow  = InstalledAppFlow.from_client_secrets_file(oauth_path, GDRIVE_OAUTH_SCOPES)
            creds = flow.run_local_server(port=0)

        # Save token for next run
        with open(token_path, "w") as f:
            f.write(creds.to_json())

    return build("drive", "v3", credentials=creds)


def get_sheets_service():
    scopes = ["https://www.googleapis.com/auth/spreadsheets"]
    creds  = ServiceAccountCredentials.from_service_account_file(GOOGLE_CREDS_FILE, scopes=scopes)
    return build("sheets", "v4", credentials=creds)


def log_job_to_sheet(job, score, assessment, tab_name="Applications", drive_link=None):
    """
    Logs a job to the specified Google Sheet tab.
    Reads the header row to find correct column indices, so it's robust to column reordering.
    Fills: Company, Position Title, URL, Linked In URL, Date, Location, Claude Score,
           Claude notes, and Resume Treatment (Drive URL, Applications tab only).
    """
    try:
        service = get_sheets_service()

        # Read header row to find column indices
        header_range = f"{tab_name}!1:1"
        header_response = service.spreadsheets().values().get(
            spreadsheetId=SHEET_ID,
            range=header_range
        ).execute()
        headers = header_response.get('values', [[]])[0]

        # Map field names to column indices (0-based)
        col_map = {}
        for idx, header in enumerate(headers):
            col_map[header.strip()] = idx

        # Check if we have all required columns
        required = ["Company", "Position Title", "URL", "Linked In URL", "Date", "Location", "Claude Score", "Claude notes"]
        missing = [col for col in required if col not in col_map]
        if missing:
            print(f"      ✗ Missing columns in {tab_name}: {missing}")
            return

        # Format location: "Remote" if remote, "Miami" if Miami-based
        location = job.get("location", "")
        if "remote" in location.lower():
            location_formatted = "Remote"
        elif "miami" in location.lower():
            location_formatted = "Miami"
        else:
            location_formatted = location

        # Format date as MM/DD/YYYY
        from datetime import datetime
        today = datetime.now().strftime("%m/%d/%Y")

        # Build a sparse row with empty strings for all columns, then fill our values
        num_cols = len(headers)
        row = [""] * num_cols

        # Fill only the columns we care about
        row[col_map["Company"]] = job.get("company", "")
        row[col_map["Position Title"]] = job.get("title", "")
        row[col_map["URL"]] = job.get("url", "")
        row[col_map["Linked In URL"]] = job.get("linkedin_url", "")  # May be empty
        row[col_map["Date"]] = today
        row[col_map["Location"]] = location_formatted
        row[col_map["Claude Score"]] = score
        row[col_map["Claude notes"]] = assessment

        # Write the Drive URL to "Resume Treatment" if the column exists and we have a link
        if drive_link and "Resume Treatment" in col_map:
            row[col_map["Resume Treatment"]] = drive_link

        # Append to sheet
        range_name = f"{tab_name}!A:ZZ"
        body = {"values": [row]}
        service.spreadsheets().values().append(
            spreadsheetId=SHEET_ID,
            range=range_name,
            valueInputOption="USER_ENTERED",
            body=body
        ).execute()

        print(f"      ✓ Logged to {tab_name} tab")

    except Exception as e:
        import traceback
        print(f"      ✗ Failed to log to {tab_name} tab: {e}")
        print(f"      Error details: {traceback.format_exc()}")


def upload_to_drive(filepath, filename, folder_id=None):
    """Uploads a file to the configured Google Drive folder.
    Uses OAuth 2.0 (personal account) so the file lands in your own Drive quota."""
    try:
        _folder = folder_id or _load_pipeline_config().get("gdrive_folder_id") or GDRIVE_FOLDER_ID
        service  = get_drive_service()
        metadata = {"name": filename, "parents": [_folder] if _folder else []}
        media    = MediaFileUpload(
            filepath,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        file = service.files().create(
            body=metadata,
            media_body=media,
            fields="id,webViewLink",
        ).execute()
        return file.get("webViewLink", "uploaded")
    except Exception as e:
        print(f"      Drive upload failed: {e}")
        return None


# ── JOB DESCRIPTION FETCHER ────────────────────────────────────────────────────

SESSION_FILE = os.path.join(SCRIPT_DIR, "linkedin_session.json")


def fetch_job_description(page, url):
    """Fetches the full job description text from a job URL.
    Uses the shared page object (with LinkedIn session loaded).
    Falls back to requests library if Playwright navigation fails (e.g., download trigger)."""
    try:
        # Primary: Try Playwright navigation (works for most sites)
        try:
            page.context.set_default_timeout(20000)
            page.goto(url, wait_until="domcontentloaded", timeout=20000)
            time.sleep(2)
        except Exception as e:
            # Fallback: Try requests library if Playwright fails (download trigger, etc.)
            if "Download" in str(e):
                print(f"      Playwright blocked by download, trying requests fallback...")
                try:
                    import requests
                    headers = {
                        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
                    }
                    response = requests.get(url, headers=headers, timeout=10)
                    print(f"      Requests status: {response.status_code}, Content-Type: {response.headers.get('Content-Type', 'unknown')}")
                    if response.status_code == 200 and "text/html" in response.headers.get("Content-Type", ""):
                        page.set_content(response.text)
                        time.sleep(1)
                        print(f"      ✓ Requests fallback succeeded")
                    else:
                        print(f"      ✗ Requests fallback failed: not HTML or bad status")
                        raise Exception("Requests fallback failed")
                except Exception as fallback_error:
                    print(f"      ✗ Requests fallback error: {fallback_error}")
                    raise
            else:
                raise

        # Expand "Show more" / "more" / "Read full description" accordions
        for more_sel in [
            "button.show-more-less-html__button--more",
            "button[aria-label='Click to see more description']",
            "a.show-more-less-html__button--more",
            "button:has-text('Show more')",
            "a:has-text('more')",
        ]:
            try:
                btn = page.query_selector(more_sel)
                if btn and btn.is_visible():
                    btn.click()
                    time.sleep(1)
                    break
            except Exception:
                pass

        # Extract job description — LinkedIn selectors first, then generic fallbacks
        desc = ""
        for selector in [
            # LinkedIn "About the job" section
            ".jobs-description__content",
            ".jobs-box__html-content",
            ".description__text",
            ".show-more-less-html__markup",
            "div.jobs-description",
            # Generic
            "[class*='description']",
            "article",
            "main",
        ]:
            el = page.query_selector(selector)
            if el:
                desc = el.inner_text().strip()
                if len(desc) > 200:
                    break

        return desc[:8000] if desc else ""

    except Exception as e:
        print(f"      Could not fetch job description: {e}")
        return ""


# ── CLAUDE SCORING ─────────────────────────────────────────────────────────────

def score_job(client, job, job_description, salary_minimum=0, preferred_locations=None):
    """Scores a job against Dom's background. Returns dict with score + assessment +
    hard disqualification flags for salary and location."""

    if preferred_locations is None:
        preferred_locations = ["remote", "miami"]

    salary_clause = ""
    if salary_minimum and salary_minimum > 0:
        salary_clause = f"""
SALARY DISQUALIFIER: The minimum acceptable total compensation is ${salary_minimum:,.0f}/year.
- If the job description states a maximum salary (top of range) BELOW ${salary_minimum:,.0f}, set salary_disqualify: true.
- If no salary is mentioned, or the salary meets or exceeds the threshold, set salary_disqualify: false.
"""

    loc_list = ", ".join(str(l).title() for l in preferred_locations)
    location_clause = f"""
LOCATION DISQUALIFIER: Acceptable work arrangements are: {loc_list}.
Read the full job description carefully for any location requirement hidden in the text:
- If it says the candidate MUST live near / within commuting distance of a specific office or hub
  that is NOT in {loc_list} (e.g., "must be within 50 miles of our NYC office", "required to be
  on-site in San Francisco"), set location_disqualify: true.
- If it says "occasional travel" (roughly once a month or less), or lists office hubs as optional
  / for those who prefer in-person, set location_disqualify: false — that is acceptable.
- If fully remote with no proximity requirement, set location_disqualify: false.
"""

    prompt = f"""{BACKGROUND_PROMPT}

---

JOB TO EVALUATE:
Title: {job.get('title')}
Company: {job.get('company')}
Location: {job.get('location')}
URL: {job.get('url')}

JOB DESCRIPTION:
{job_description if job_description else "[Could not fetch — use title/company to infer]"}

---
{salary_clause}{location_clause}
---

Score this role for Dom on a scale of 0–10 based on the scoring criteria above. Use the "Top scoring roles from prior sessions" as calibration — compare this role to those examples.

IMPORTANT: Prioritize role-type alignment (30% weight) over industry fit (20% weight). If the role type is a direct match (e.g., agentic AI role + Dom's agent framework, or protocol PM + EON/L3 experience), score highly even if the industry is outside crypto/fintech. Industry matters less when the actual work is a perfect fit.

Be direct and honest. If it's a poor fit, say so clearly (e.g., "hard pass", "would be doing you a disservice"). If there are hard gaps — specific must-have requirements Dom cannot claim — list them as bullets under "Hard gaps:" in your assessment.

Structure your assessment like this:
- Opening line with frank verdict (e.g., "Strong fit", "Worth applying despite X", "Hard pass")
- If there are specific hard gaps (domain expertise, certifications, industry background Dom doesn't have), list them as bullets under "Hard gaps:"
- 1-2 lines explaining why the gaps matter or why the fit works
- Flag compensation if it's notably low (under $160K)

Return ONLY a JSON object:
{{
  "score": 8.5,
  "assessment": "Your full assessment here with hard gaps as bullets if applicable",
  "role_type": "blockchain|fintech|ai|consumer|compliance|marketplace|other",
  "title_for_file": "Senior_PM_Crypto",
  "salary_disqualify": false,
  "location_disqualify": false,
  "disqualify_reason": ""
}}

No other text. Just the JSON."""

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1200,
        messages=[{"role": "user", "content": prompt}]
    )

    raw = response.content[0].text.strip()
    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'^```\s*',     '', raw)
    raw = re.sub(r'\s*```$',     '', raw)

    try:
        return json.loads(raw)
    except Exception:
        return {"score": 0, "assessment": "Could not parse score", "role_type": "other", "title_for_file": "Senior_PM"}


# ── RESUME BUILDER ─────────────────────────────────────────────────────────────

def build_resume_content(client, job, job_description, score_data):
    """Asks Claude to produce the full tailored resume content as structured JSON."""

    role_type   = score_data.get("role_type", "other")
    score       = score_data.get("score", 0)
    assessment  = score_data.get("assessment", "")

    resume_path = None
    base_resume_dir = os.path.join(SCRIPT_DIR, "base_resume")
    for ext in (".pdf", ".docx"):
        import glob as _glob
        matches = _glob.glob(os.path.join(base_resume_dir, f"*{ext}"))
        if matches:
            resume_path = matches[0]
            break
    if resume_path is None:
        raise FileNotFoundError(
            "No base resume found in base_resume/. "
            "Upload your resume (.pdf or .docx) via the web UI."
        )

    if resume_path.endswith(".pdf"):
        with open(resume_path, "rb") as f:
            resume_content_block = {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": base64.standard_b64encode(f.read()).decode("utf-8"),
                },
            }
    else:
        from docx import Document
        doc = Document(resume_path)
        resume_text = "\n".join(p.text for p in doc.paragraphs)
        resume_content_block = {"type": "text", "text": resume_text}

    text_prompt = f"""{BACKGROUND_PROMPT}

═══════════════════════════════════════════════════
YOUR TASK: BUILD A TAILORED RESUME
═══════════════════════════════════════════════════
Role:       {job.get('title')}
Company:    {job.get('company')}
Location:   {job.get('location')}
Role type:  {role_type}
Score:      {score}/10
Assessment: {assessment}

JOB DESCRIPTION:
{job_description if job_description else "[Use title/company/role type to tailor]"}

═══════════════════════════════════════════════════
BEFORE WRITING: DO THESE THREE STEPS INTERNALLY
═══════════════════════════════════════════════════
STEP 1 — IDENTIFY THE TOP 3 SIGNALS
  What are the 3 strongest signals in Dom's background for this specific role?
  These are not the most impressive things overall — they are the most
  relevant to what this hiring team is actually testing for.

STEP 2 — IDENTIFY THE HONEST GAPS
  What 1-2 requirements in this JD does Dom genuinely not have?
  Do not paper over them. Do not claim them. Work around them.

STEP 3 — DECIDE THE LEAD BULLET
  The first bullet of the most relevant role block must be the single best
  signal for this hiring team. Not the Vela bullet by default. Not the $200M
  migration by default. The right bullet for THIS role.

═══════════════════════════════════════════════════
OUTPUT FORMAT
═══════════════════════════════════════════════════
Return ONLY a JSON object. No markdown. No extra text. No explanation.

{{
  "name": "YOUR FULL NAME",
  "contact": "City, State  •  Phone  •  Email  •  LinkedIn  •  Work Authorization",
  "linkedin_url": "https://www.linkedin.com/in/yourprofile/",
  "summary": "Strategic Product Leader with 15+ years of experience...",
  "competencies": [
    {{ "category": "Category Name", "skills": "Skill 1, Skill 2, Skill 3..." }}
  ],
  "experience": [
    {{
      "title": "Senior Product Manager",
      "company": "Horizen Labs",
      "location": "Remote",
      "dates": "Jan 2024 – Mar 2026",
      "bullets": [
        "First bullet — most relevant signal for this role...",
        "Second bullet..."
      ]
    }}
  ],
  "entrepreneurship": {{
    "header": "Founder, Product & Technical Lead | Disci.io | Miami, FL  Ongoing",
    "bullets": [
      "Sole technical founder...",
      "Second bullet..."
    ]
  }},
  "education": [
    "MIT Sloan School of Management — Blockchain: Business Innovation & Application (2021)",
    "Queens College, CUNY — B.A., Computer Science"
  ]
}}"""

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=6000,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": "This is Dom's existing resume. Use the bullet points and wording from this document as your source of truth. Reorder bullets to lead with the most relevant signal for this role. Reframe where needed for the JD. Do not invent new bullets, do not expand existing bullets beyond what is written here, and do not add claims that are not already present in this document.",
                },
                resume_content_block,
                {
                    "type": "text",
                    "text": text_prompt,
                },
            ],
        }]
    )

    raw = response.content[0].text.strip()
    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'^```\s*',     '', raw)
    raw = re.sub(r'\s*```$',     '', raw)

    try:
        return json.loads(raw)
    except Exception as e:
        print(f"   ⚠️  Could not parse resume JSON: {e}")
        return None


def build_docx(resume_data, output_path):
    """Generates a .docx file from resume JSON using docx-js via Node."""

    js_code = r"""
const { Document, Packer, Paragraph, TextRun, AlignmentType,
        LevelFormat, ExternalHyperlink, BorderStyle, HeadingLevel,
        UnderlineType } = require('docx');
const fs = require('fs');

const data = JSON.parse(fs.readFileSync(process.argv[3], 'utf8'));

const COLORS = {
  navy: '1F4E79',    // name, section headings, borders
  black: '000000',   // body text, bullets
  gray: '595959',    // contact line, role metadata, skill lists
  link: '1155CC',    // LinkedIn hyperlink
};
const FONT = 'Arial';

function sectionHeading(text) {
  return new Paragraph({
    children: [new TextRun({ text: text.toUpperCase(), bold: true, size: 22, font: FONT, color: COLORS.navy })],
    border: { bottom: { style: BorderStyle.SINGLE, size: 6, color: COLORS.navy, space: 1 } },
    spacing: { before: 140, after: 40 }
  });
}

function bullet(text) {
  return new Paragraph({
    numbering: { reference: 'bullets', level: 0 },
    children: [new TextRun({ text, size: 20, font: FONT, color: COLORS.black })],
    spacing: { before: 20, after: 20 }
  });
}

// Role header: bold black title  |  gray company  |  gray location  |  gray dates
function roleHeader(title, company, location, dates) {
  const meta = [company, location, dates].filter(s => s).join('  |  ');
  return new Paragraph({
    children: [
      new TextRun({ text: title, bold: true, size: 20, font: FONT, color: COLORS.black }),
      new TextRun({ text: '  |  ' + meta, size: 20, font: FONT, color: COLORS.gray }),
    ],
    spacing: { before: 120, after: 30 }
  });
}

const children = [];

// ── NAME ──
children.push(new Paragraph({
  alignment: AlignmentType.CENTER,
  children: [new TextRun({ text: data.name, bold: true, size: 32, font: FONT, color: COLORS.navy })],
  spacing: { before: 0, after: 40 }
}));

// ── CONTACT LINE with hyperlinked LinkedIn ──
const contactParts = data.contact.split('LinkedIn');
children.push(new Paragraph({
  alignment: AlignmentType.CENTER,
  children: [
    new TextRun({ text: contactParts[0], size: 18, font: FONT, color: COLORS.gray }),
    new ExternalHyperlink({
      link: data.linkedin_url,
      children: [new TextRun({ text: 'LinkedIn', size: 18, font: FONT, color: COLORS.link, underline: { type: UnderlineType.SINGLE } })]
    }),
    new TextRun({ text: contactParts[1] || '', size: 18, font: FONT, color: COLORS.gray }),
  ],
  spacing: { before: 0, after: 60 }
}));

// ── SUMMARY ──
children.push(sectionHeading('Professional Summary'));
children.push(new Paragraph({
  children: [new TextRun({ text: data.summary, size: 20, font: FONT, color: COLORS.black })],
  spacing: { before: 60, after: 80 }
}));

// ── CORE COMPETENCIES ──
children.push(sectionHeading('Core Competencies'));
data.competencies.forEach(c => {
  children.push(new Paragraph({
    children: [
      new TextRun({ text: c.category + ': ', bold: true, size: 19, font: FONT, color: COLORS.black }),
      new TextRun({ text: c.skills, size: 19, font: FONT, color: COLORS.gray })
    ],
    spacing: { before: 30, after: 30 }
  }));
});

// ── EXPERIENCE ──
children.push(sectionHeading('Professional Experience'));
data.experience.forEach(role => {
  children.push(roleHeader(role.title, role.company, role.location, role.dates));
  role.bullets.forEach(b => children.push(bullet(b)));
});

// ── ENTREPRENEURSHIP ──
if (data.entrepreneurship) {
  children.push(sectionHeading('Entrepreneurship'));
  // Parse header: expect "Title | Company | Location | Dates" or just a string
  const hdr = data.entrepreneurship.header || '';
  const hdrParts = hdr.split('|').map(s => s.trim());
  if (hdrParts.length >= 4) {
    children.push(roleHeader(hdrParts[0], hdrParts[1], hdrParts[2], hdrParts[3]));
  } else if (hdrParts.length === 3) {
    // 3 parts: "Title | Company | Location Dates"
    children.push(roleHeader(hdrParts[0], hdrParts[1], hdrParts[2], ''));
  } else if (hdrParts.length === 2) {
    children.push(new Paragraph({
      children: [
        new TextRun({ text: hdrParts[0], bold: true, size: 20, font: FONT, color: COLORS.black }),
        new TextRun({ text: '  |  ' + hdrParts[1], size: 20, font: FONT, color: COLORS.gray }),
      ],
      spacing: { before: 80, after: 30 }
    }));
  } else {
    // Fallback: bold title, rest gray
    children.push(new Paragraph({
      children: [new TextRun({ text: hdr, bold: true, size: 20, font: FONT, color: COLORS.black })],
      spacing: { before: 80, after: 30 }
    }));
  }
  data.entrepreneurship.bullets.forEach(b => children.push(bullet(b)));
}

// ── EDUCATION ──
children.push(sectionHeading('Education & Certifications'));
data.education.forEach((e, i) => {
  children.push(new Paragraph({
    children: [new TextRun({ text: e, size: 20, font: FONT, color: COLORS.black })],
    spacing: { before: i === 0 ? 60 : 0, after: i === 0 ? 20 : 0 }
  }));
});

// ── BUILD DOC ──
const doc = new Document({
  numbering: {
    config: [{
      reference: 'bullets',
      levels: [{ level: 0, format: LevelFormat.BULLET, text: '\u2022', alignment: AlignmentType.LEFT,
        style: { paragraph: { indent: { left: 480, hanging: 240 } },
          run: { font: FONT, size: 20, color: COLORS.black } } }]
    }]
  },
  styles: {
    default: { document: { run: { font: FONT, size: 20, color: COLORS.black } } }
  },
  sections: [{
    properties: {
      page: {
        size: { width: 12240, height: 15840 },
        margin: { top: 720, right: 900, bottom: 720, left: 900 }
      }
    },
    children
  }]
});

Packer.toBuffer(doc).then(buf => {
  fs.writeFileSync(process.argv[2], buf);
  console.log('ok');
});
"""

    with tempfile.TemporaryDirectory() as tmp:
        # Write resume data JSON
        data_path = os.path.join(tmp, "resume_data.json")
        with open(data_path, "w") as f:
            json.dump(resume_data, f)

        # Write JS builder
        js_path = os.path.join(tmp, "build.js")
        with open(js_path, "w") as f:
            f.write(js_code)

        # Run Node from the PROJECT directory so require('docx') resolves
        # to the project's node_modules, while reading data from the temp dir
        result = subprocess.run(
            ["node", js_path, output_path, data_path],
            cwd=SCRIPT_DIR,
            env={**os.environ, "NODE_PATH": os.path.join(SCRIPT_DIR, "node_modules")},
            capture_output=True, text=True
        )

        if result.returncode != 0 or "ok" not in result.stdout:
            print(f"      docx build error: {result.stderr[:300]}")
            return False

        return True


# ── MAIN PIPELINE ──────────────────────────────────────────────────────────────

def run_pipeline(jobs, test_scoring_only=False):
    """
    Main pipeline. Accepts a list of job dicts with keys: title, company, location, url.
    Scores each, builds resumes for score >= score_threshold, uploads to Drive.

    Args:
        jobs: List of job dicts
        test_scoring_only: If True, skips resume building and only tests scoring + sheet logging
    """
    # Re-read config at run time so settings changes take effect without restart
    pipeline_cfg       = _load_pipeline_config()
    score_threshold    = pipeline_cfg["score_threshold"]
    salary_minimum     = pipeline_cfg["salary_minimum"]
    preferred_locs     = pipeline_cfg["preferred_locations"]
    gdrive_folder      = pipeline_cfg.get("gdrive_folder_id") or GDRIVE_FOLDER_ID

    print(f"\n{'='*65}")
    mode = "TEST MODE - Scoring Only" if test_scoring_only else "Starting"
    print(f"  RESUME PIPELINE -- {mode}")
    print(f"  Processing {len(jobs)} jobs  |  Score threshold: {score_threshold}/10")
    if salary_minimum:
        print(f"  Salary minimum:    ${salary_minimum:,.0f}/yr  (from Settings)")
    print(f"  Location filter:   {preferred_locs}  (from Settings)")
    if test_scoring_only:
        print("  (Resume building DISABLED for testing)")
    print(f"{'='*65}\n")

    client   = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    results  = []
    built    = 0

    # Open one browser session with LinkedIn cookies for all JD fetching
    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=True,
            args=[
                "--no-sandbox",
                "--disable-blink-features=AutomationControlled"
            ]
        )

        # Create context with realistic browser fingerprint to avoid bot detection
        context_options = {
            "viewport": {"width": 1920, "height": 1080},
            "user_agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
        }

        if os.path.exists(SESSION_FILE):
            context = browser.new_context(storage_state=SESSION_FILE, **context_options)
        else:
            context = browser.new_context(**context_options)

        page = context.new_page()

        for i, job in enumerate(jobs, 1):
            title   = job.get("title", "Unknown")
            company = job.get("company", "Unknown")
            url     = job.get("url", "")
            linkedin_url = job.get("linkedin_url", "")

            print(f"[{i}/{len(jobs)}]  {title} @ {company}")

            # 1. Fetch job description
            #    Primary:    apply URL (always the authoritative source)
            #    Fallback 1: pre-scraped listing text (BuiltIn "The Role")
            #    Fallback 2: listing page URL (LinkedIn "About the job" / BuiltIn page)
            #    No JD:      log to Applications and skip — don't create an application
            print("      Fetching job description...")
            jd = ""

            if url:
                jd = fetch_job_description(page, url)
                if jd:
                    print(f"      Got description from apply URL ({len(jd)} chars)")

            if not jd:
                scraped_desc = job.get("description", "")
                if scraped_desc and len(scraped_desc) > 200:
                    jd = scraped_desc
                    print(f"      Using pre-scraped listing description ({len(jd)} chars)")

            if not jd and linkedin_url and linkedin_url != url:
                jd = fetch_job_description(page, linkedin_url)
                if jd:
                    print(f"      Got description from listing page ({len(jd)} chars)")

            if not jd:
                print("      Could not fetch job description by any means — logging and skipping")
                log_job_to_sheet(
                    job, "N/A",
                    "Could not fetch job description — no application created",
                    tab_name="Applications",
                )
                continue

            # 2. Score + hard disqualification checks
            print("      Scoring with Claude...")
            score_data        = score_job(client, job, jd, salary_minimum, preferred_locs)
            score             = score_data.get("score", 0)
            assessment        = score_data.get("assessment", "")
            salary_disqualify = score_data.get("salary_disqualify", False)
            loc_disqualify    = score_data.get("location_disqualify", False)
            disq_reason       = score_data.get("disqualify_reason", "")

            print(f"      Score: {score}/10 -- {assessment}")

            # Hard disqualifiers → Skips tab regardless of score
            if salary_disqualify or loc_disqualify:
                tag = "Salary below minimum" if salary_disqualify else "Location mismatch"
                reason = disq_reason or tag
                print(f"      ✗ DISQUALIFIED ({tag}): {reason}")
                log_job_to_sheet(job, score, f"[DISQUALIFIED — {tag}] {assessment}", tab_name="Skips")
                continue

            result = {**job, "score": score, "assessment": assessment, "resume_built": False, "drive_link": None}

            # 3. Build resume if score meets threshold (unless in test mode)
            if score >= score_threshold:
                if test_scoring_only:
                    print(f"      Score >= {score_threshold} -- would build resume (skipped in test mode)")
                    log_job_to_sheet(job, score, assessment, tab_name="Applications")
                else:
                    print(f"      Score >= {score_threshold} -- building resume...")

                    resume_data = build_resume_content(client, job, jd, score_data)
                    drive_link  = None

                    if resume_data:
                        _resume_prefix = pipeline_cfg.get("resume_prefix", "My_Resume")
                        title_slug  = score_data.get("title_for_file", "Senior_PM").replace(" ", "_")
                        # Only append company if it's not already in the slug
                        company_slug = company.replace(" ", "_")
                        if company_slug.lower() not in title_slug.lower():
                            filename = f"{_resume_prefix}_{title_slug}_{company_slug}.docx"
                        else:
                            filename = f"{_resume_prefix}_{title_slug}.docx"
                        # Save to a resumes/ folder in the project directory
                        resumes_dir = os.path.join(SCRIPT_DIR, "resumes")
                        os.makedirs(resumes_dir, exist_ok=True)
                        output_path = os.path.join(resumes_dir, filename)

                        success = build_docx(resume_data, output_path)

                        if success:
                            print(f"      Uploading to Google Drive...")
                            drive_link = upload_to_drive(output_path, filename, folder_id=gdrive_folder)
                            result["resume_built"] = True
                            result["drive_link"]   = drive_link
                            result["filename"]     = filename
                            built += 1
                            if drive_link:
                                print(f"      Drive link: {drive_link}")
                        else:
                            print("      Resume build failed")

                    # Log to Applications tab with the Drive link (if the build succeeded)
                    log_job_to_sheet(job, score, assessment, tab_name="Applications", drive_link=drive_link)
            else:
                print(f"      Score below {score_threshold} -- skipping")
                # Log to Skips tab
                print(f"      Attempting to log to Skips tab...")
                log_job_to_sheet(job, score, assessment, tab_name="Skips")

            results.append(result)
            print()
            time.sleep(1)

        browser.close()

    # ── SUMMARY ──
    print(f"\n{'='*65}")
    print(f"  PIPELINE COMPLETE -- {built} resumes built and uploaded")
    print(f"{'='*65}")
    print(f"\n  {'ROLE':<45} {'SCORE':>6}  {'RESUME'}")
    print(f"  {'-'*58}")
    for r in sorted(results, key=lambda x: x['score'] if isinstance(x['score'], (int, float)) else -1, reverse=True):
        status = f"OK {r.get('filename','')}" if r['resume_built'] else "--"
        label  = f"{r['title'][:30]} @ {r['company'][:14]}"
        score_str = f"{r['score']:>5.1f}" if isinstance(r['score'], (int, float)) else f"{r['score']:>5}"
        print(f"  {label:<45} {score_str}  {status}")
    print()


# ── ENTRY POINT ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    # Check for flags
    test_mode = "--test-scoring" in sys.argv
    if test_mode:
        sys.argv.remove("--test-scoring")

    # Check for a JSON file passed as argument (from job_scraper integration)
    if len(sys.argv) > 1 and os.path.exists(sys.argv[1]):
        with open(sys.argv[1]) as f:
            jobs = json.load(f)
        print(f"Loaded {len(jobs)} jobs from {sys.argv[1]}")
        run_pipeline(jobs, test_scoring_only=test_mode)
    else:
        print("Usage: python3 resume_pipeline.py <jobs.json> [--test-scoring]")
        print("       or import run_pipeline() from job_scraper.py")
        print()
        print("The jobs.json file should be an array of objects with:")
        print('  title, company, location, url, linkedin_url')
        print()
        print("Options:")
        print("  --test-scoring  Skip resume building, only test scoring and sheet logging")
